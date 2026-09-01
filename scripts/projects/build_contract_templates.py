"""Build sanitized project document templates from user-owned source files.

The resulting towage DOCX contains visible placeholder tokens consumed by the
SeaPilot browser generator. The BIMCO PDF contains only the generic Part II
pages from the supplied executed contract; signed/customer-specific pages are
deliberately excluded. The bareboat-charter DOCX keeps the supplied contract
layout and clauses while removing every executed value and signature.
"""

from __future__ import annotations

import argparse
from copy import deepcopy
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from docx import Document
from docx.table import _Cell
from lxml import etree
from pypdf import PdfReader, PdfWriter


WORD_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
REL_NS = 'http://schemas.openxmlformats.org/package/2006/relationships'
NS = {'w': WORD_NS}


def set_cell(cell: _Cell, value: str) -> None:
    for drawing in cell._element.xpath('.//w:drawing'):
        drawing.getparent().remove(drawing)
    cell.text = value


def build_towage_template(source: Path, target: Path) -> None:
    document = Document(source)
    clauses = document.tables[0]

    replacements = {
        (1, 1): '{{CONTRACT_DATE_LONG}}',
        (3, 0): '{{CHARTERER}}',
        (3, 1): '{{OWNER}}',
        (5, 0): '{{TOWED_VESSEL}}',
        (5, 1): '{{TUG}}',
        (7, 0): '{{TOWED_CONDITIONS}}',
        (9, 0): '{{PICKUP_PLACE}}',
        (9, 1): '{{DEPARTURE_WINDOW}}',
        (11, 0): '{{DESTINATION_PLACE}}',
        (11, 1): '{{ARRIVAL_WINDOW}}',
        (13, 0): '{{CONNECTION_TIME}}',
        (13, 1): '{{DISCONNECTION_TIME}}',
        (15, 0): '{{FIXED_PRICE}}',
        (15, 1): '{{OPTIONAL_COSTS}}',
        (17, 0): '{{PAYMENT_TERMS}}',
        (17, 1): '{{ADDITIONAL_CHARGES}}',
        (19, 0): '{{SPECIAL_CONDITIONS}}',
        (21, 0): '{{CHARTERER_SIGNATORY}}\n{{SIGNATURE_DATE}}',
        (21, 1): '{{OWNER_SIGNATORY}}\n{{SIGNATURE_DATE}}',
    }
    for (row, column), value in replacements.items():
        set_cell(clauses.cell(row, column), value)

    signatures = document.tables[1]
    set_cell(signatures.cell(1, 0), '{{OWNER_SIGNATORY}}\n{{SIGNATURE_DATE}}')
    set_cell(signatures.cell(1, 1), '{{CHARTERER_SIGNATORY}}\n{{SIGNATURE_DATE}}')

    processed_headers: set[int] = set()
    for section in document.sections:
        for table in section.header.tables:
            element_id = id(table._element)
            if element_id in processed_headers:
                continue
            processed_headers.add(element_id)
            set_cell(table.cell(0, 6), '{{CONTRACT_DATE_SHORT}}')
            set_cell(table.cell(1, 1), '{{DOCUMENT_CODE}}')
            set_cell(table.cell(1, 2), '{{PROJECT_CODE}}')

    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)


def build_supplytime_part_ii(source: Path, target: Path) -> None:
    reader = PdfReader(source)
    if reader.is_encrypted:
        reader.decrypt('')
    if len(reader.pages) < 24:
        raise ValueError('The supplied SUPPLYTIME document does not contain the expected Part II pages.')

    writer = PdfWriter()
    # Source pages 5–24 are the unfilled, generic SUPPLYTIME 2017 Part II clauses.
    # Pages 1–4 and 25–30 contain executed contract data and are never copied.
    for page in reader.pages[4:24]:
        writer.add_page(page)
    writer.add_metadata({
        '/Title': 'SUPPLYTIME 2017 - Part II',
        '/Author': 'BIMCO',
        '/Subject': 'SeaPilot internal authorized contract template',
    })
    target.parent.mkdir(parents=True, exist_ok=True)
    with target.open('wb') as output:
        writer.write(output)


def _word_tag(local_name: str) -> str:
    return f'{{{WORD_NS}}}{local_name}'


def _replace_paragraph_content(paragraph: etree._Element, value: str) -> None:
    paragraph_properties = paragraph.find('w:pPr', NS)
    first_run = paragraph.find('w:r', NS)
    run_properties = first_run.find('w:rPr', NS) if first_run is not None else None
    for child in list(paragraph):
        if child.tag != _word_tag('pPr'):
            paragraph.remove(child)
    if value:
        run = etree.SubElement(paragraph, _word_tag('r'))
        if run_properties is not None:
            run.append(deepcopy(run_properties))
        for index, line in enumerate(value.split('\n')):
            if index:
                etree.SubElement(run, _word_tag('br'))
            text = etree.SubElement(run, _word_tag('t'))
            if line.startswith(' ') or line.endswith(' '):
                text.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            text.text = line
    elif paragraph_properties is None:
        etree.SubElement(paragraph, _word_tag('pPr'))


def _replace_cell_content(cell: etree._Element, value: str) -> None:
    first_paragraph = cell.find('w:p', NS)
    paragraph_properties = (
        deepcopy(first_paragraph.find('w:pPr', NS))
        if first_paragraph is not None and first_paragraph.find('w:pPr', NS) is not None
        else None
    )
    first_run = first_paragraph.find('w:r', NS) if first_paragraph is not None else None
    run_properties = (
        deepcopy(first_run.find('w:rPr', NS))
        if first_run is not None and first_run.find('w:rPr', NS) is not None
        else None
    )
    for child in list(cell):
        if child.tag != _word_tag('tcPr'):
            cell.remove(child)
    paragraph = etree.SubElement(cell, _word_tag('p'))
    if paragraph_properties is not None:
        paragraph.append(paragraph_properties)
    if not value:
        return
    run = etree.SubElement(paragraph, _word_tag('r'))
    if run_properties is not None:
        run.append(run_properties)
    for index, line in enumerate(value.split('\n')):
        if index:
            etree.SubElement(run, _word_tag('br'))
        text = etree.SubElement(run, _word_tag('t'))
        if line.startswith(' ') or line.endswith(' '):
            text.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        text.text = line


def _table_cells(table: etree._Element, row_index: int) -> list[etree._Element]:
    rows = table.findall('w:tr', NS)
    return rows[row_index].findall('w:tc', NS)


def _set_row_minimum_height(table: etree._Element, row_index: int, twips: int) -> None:
    row = table.findall('w:tr', NS)[row_index]
    properties = row.find('w:trPr', NS)
    if properties is None:
        properties = etree.Element(_word_tag('trPr'))
        row.insert(0, properties)
    height = properties.find('w:trHeight', NS)
    if height is None:
        height = etree.SubElement(properties, _word_tag('trHeight'))
    height.set(_word_tag('val'), str(twips))
    height.set(_word_tag('hRule'), 'atLeast')


def _without_executed_signature_relationships(
    relationships_xml: bytes,
) -> tuple[bytes, set[str]]:
    root = etree.fromstring(relationships_xml)
    removed_targets: set[str] = set()
    for relationship in list(root):
        if relationship.get('Target') == 'media/image1.jpeg':
            removed_targets.add(relationship.get('Target', ''))
            root.remove(relationship)
    return etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True), removed_targets


def build_bareboat_template(source: Path, target: Path) -> None:
    """Remove executed values while preserving the reference layout and clauses."""

    with ZipFile(source) as archive:
        parts = {info.filename: (info, archive.read(info.filename)) for info in archive.infolist()}

    document_root = etree.fromstring(parts['word/document.xml'][1])
    header_root = etree.fromstring(parts['word/header1.xml'][1])

    body_tables = document_root.findall('.//w:body/w:tbl', NS)
    if len(body_tables) < 2:
        raise ValueError('The bareboat-charter source does not contain the expected two body tables.')

    main_table = body_tables[0]
    for row_index in (1, 3, 5, 7, 9, 11, 13, 15, 17, 19, 21):
        for cell in _table_cells(main_table, row_index):
            _replace_cell_content(cell, '')
    # These three source rows were auto-sized by their executed values. Keep
    # the measured source geometry after removing the sensitive text.
    _set_row_minimum_height(main_table, 5, 1356)
    _set_row_minimum_height(main_table, 7, 552)
    _set_row_minimum_height(main_table, 9, 816)

    signature_table = body_tables[1]
    signature_header_cells = _table_cells(signature_table, 0)
    if len(signature_header_cells) != 2:
        raise ValueError('The bareboat-charter signature table does not contain two columns.')
    _replace_cell_content(signature_header_cells[0], 'Pour l’affréteur (2) :')
    _replace_cell_content(signature_header_cells[1], 'Pour le propriétaire (3) :')
    for cell in _table_cells(signature_table, 1):
        _replace_cell_content(cell, '')

    for paragraph in document_root.findall('.//w:body/w:p', NS):
        text = ''.join(paragraph.itertext()).strip()
        if text.startswith('Fait à '):
            _replace_paragraph_content(paragraph, '')

    header_tables = header_root.findall('.//w:tbl', NS)
    if len(header_tables) != 1:
        raise ValueError('The bareboat-charter source does not contain the expected header table.')
    header_table = header_tables[0]
    date_cells = _table_cells(header_table, 0)
    identity_cells = _table_cells(header_table, 1)
    if len(date_cells) != 7 or len(identity_cells) != 4:
        raise ValueError('The bareboat-charter header geometry differs from the supplied reference.')
    _replace_cell_content(date_cells[4], '')
    _replace_cell_content(identity_cells[2], '')
    _replace_cell_content(identity_cells[3], 'CONTRAT D’AFFRETEMENT COQUE NUE')

    parts['word/document.xml'] = (
        parts['word/document.xml'][0],
        etree.tostring(document_root, xml_declaration=True, encoding='UTF-8', standalone=True),
    )
    parts['word/header1.xml'] = (
        parts['word/header1.xml'][0],
        etree.tostring(header_root, xml_declaration=True, encoding='UTF-8', standalone=True),
    )
    relationships, removed_targets = _without_executed_signature_relationships(
        parts['word/_rels/document.xml.rels'][1],
    )
    parts['word/_rels/document.xml.rels'] = (
        parts['word/_rels/document.xml.rels'][0],
        relationships,
    )
    for removed_target in removed_targets:
        parts.pop(f'word/{removed_target}', None)

    target.parent.mkdir(parents=True, exist_ok=True)
    with ZipFile(target, 'w', compression=ZIP_DEFLATED, compresslevel=9) as archive:
        for name, (source_info, payload) in parts.items():
            info = ZipInfo(filename=name, date_time=source_info.date_time)
            info.compress_type = ZIP_DEFLATED
            info.external_attr = source_info.external_attr
            info.create_system = source_info.create_system
            archive.writestr(info, payload)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument('--towage-source', type=Path)
    parser.add_argument('--bimco-source', type=Path)
    parser.add_argument('--bareboat-source', type=Path)
    parser.add_argument('--output-dir', type=Path, required=True)
    args = parser.parse_args()

    if not any((args.towage_source, args.bimco_source, args.bareboat_source)):
        parser.error('Provide at least one source document.')
    if args.towage_source:
        build_towage_template(args.towage_source, args.output_dir / 'contrat-remorquage-bbtm.docx')
    if args.bimco_source:
        build_supplytime_part_ii(args.bimco_source, args.output_dir / 'supplytime-2017-part-ii.pdf')
    if args.bareboat_source:
        build_bareboat_template(args.bareboat_source, args.output_dir / 'contrat-affretement-bbtm.docx')


if __name__ == '__main__':
    main()
